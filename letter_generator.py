"""
Generation de lettres Word (.docx) personnalisees a partir du template Mirabaud.
Copie le template et remplace les champs personnalisables (date, dirigeant, entreprise, secteur).
Tout le reste (header Mirabaud, footer legal, mise en page) est preserve.
"""

import os
import re
import requests
from io import BytesIO
from datetime import datetime
from urllib.parse import urlparse

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


MOIS_FR = {
    'January': 'janvier', 'February': 'février', 'March': 'mars',
    'April': 'avril', 'May': 'mai', 'June': 'juin',
    'July': 'juillet', 'August': 'août', 'September': 'septembre',
    'October': 'octobre', 'November': 'novembre', 'December': 'décembre'
}

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'templates', 'lettre_template.docx')

PRENOMS_FEMININS = {
    'maria', 'marie', 'anne', 'sophie', 'catherine', 'isabelle', 'nathalie',
    'christine', 'sylvie', 'patricia', 'florence', 'sandrine', 'valérie',
    'valerie', 'caroline', 'julie', 'laura', 'sarah', 'emma', 'alice',
    'claire', 'marguerite', 'louise', 'charlotte', 'camille', 'lucie',
    'brigitte', 'monique', 'nicole', 'danielle', 'françoise', 'francoise',
    'martine', 'véronique', 'veronique', 'dominique', 'céline', 'celine',
    'audrey', 'stéphanie', 'stephanie', 'virginie', 'delphine', 'hélène',
    'helene', 'emilie', 'alexandra', 'béatrice', 'beatrice', 'corinne',
    'myriam', 'muriel', 'chantal', 'agnès', 'agnes', 'laurence',
    'annick', 'joëlle', 'joelle', 'michèle', 'michele', 'pascale',
    'rosalie', 'sabine', 'ségolène', 'segolene', 'élise', 'elise',
    'laetitia', 'léa', 'lea', 'manon', 'océane', 'oceane', 'pauline',
    'amandine', 'mathilde', 'juliette', 'clémence', 'clemence', 'eva',
    'inès', 'ines', 'agathe', 'constance', 'gaëlle', 'gaelle',
    'madeleine', 'jeanne', 'thérèse', 'therese', 'simone', 'yvonne',
    'geneviève', 'genevieve', 'colette', 'jacqueline', 'josette',
    'mireille', 'odette', 'suzanne', 'antoinette', 'bernadette',
    'denise', 'germaine', 'henriette', 'lucienne', 'marcelle',
    'pierrette', 'renée', 'renee', 'solange', 'yvette',
    'melissa', 'mélissa', 'nadia', 'sonia', 'diana', 'lina',
    'elena', 'élena', 'marina', 'nina', 'rosa', 'lola',
    'clara', 'sara', 'anna', 'hannah', 'jessica', 'jennifer',
    'christina', 'nathalia', 'tatiana', 'victoria', 'patricia',
    'helena', 'hélèna', 'fatima', 'karima', 'samira', 'yasmina',
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _clean(val, default=''):
    """Nettoie une valeur qui peut etre None, NaN, ou valide."""
    if val is None:
        return default
    if isinstance(val, float) and val != val:  # NaN check
        return default
    return str(val) if val else default


def _format_date_fr():
    """Retourne la date du jour en francais : '24 février 2026'"""
    date_str = datetime.now().strftime("%d %B %Y")
    for en, fr in MOIS_FR.items():
        date_str = date_str.replace(en, fr)
    if date_str.startswith('0'):
        date_str = date_str[1:]
    return date_str


def _extract_domain(site_web):
    """Extrait le domaine d'une URL de site web."""
    if not site_web:
        return ''
    url = str(site_web).strip()
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    try:
        return urlparse(url).netloc.lower().replace('www.', '')
    except Exception:
        return ''


def _clean_company_name(name):
    """
    Nettoie le nom d'entreprise :
    - 'HERVE (HERVE)' -> 'Herve'
    - 'CCF (CCF - BANQUE DES CARAIBES)' -> 'CCF'
    - 'SOCIETE GENERALE (SG)' -> 'Société Générale'
    - 'ELIOR RESTAURATION FRANCE (ELIOR...)' -> 'Elior Restauration France'
    """
    if not name:
        return name

    # Doublon exact : "HERVE (HERVE)" -> "Herve"
    m = re.match(r'^(.+?)\s*\(\1\)\s*$', name, re.IGNORECASE)
    if m:
        return m.group(1).strip().title()

    # Nom court suivi d'une parenthese qui commence pareil : "CCF (CCF - BANQUE...)" -> "CCF"
    m = re.match(r'^([A-Z\s\-]{2,10}?)\s*\(\1\b.*\)\s*$', name, re.IGNORECASE)
    if m:
        return m.group(1).strip()

    # Sigle court entre parentheses : "SOCIETE GENERALE (SG)" -> partie avant les parens
    m = re.match(r'^(.+?)\s*\([A-Z\-]{1,5}\)\s*$', name)
    if m:
        return m.group(1).strip().title()

    # Nom trop long avec parentheses -> partie avant les parens en Titre
    if len(name) > 30 and '(' in name:
        before_paren = name.split('(')[0].strip()
        if before_paren:
            return before_paren.title()

    return name


def _format_civilite(dirigeant_full, nom='', prenom=''):
    """
    Retourne la salutation formatee : "Madame Borel," ou "Monsieur Cromb,"
    Si nom et prenom sont fournis (champs separes de l'API), les utiliser directement.
    Sinon parser le string complet.
    """
    if not nom and not dirigeant_full:
        return 'Madame, Monsieur,'

    # Personne morale uniquement → salutation generique
    if dirigeant_full and dirigeant_full.startswith('PM:'):
        return 'Madame, Monsieur,'

    # Mode 1 : nom et prenom separes (depuis l'API)
    if nom:
        nom_formate = nom.strip().title()
        premier_prenom = (prenom.split()[0].split('-')[0] if prenom else '').lower()
        civilite = 'Madame' if premier_prenom in PRENOMS_FEMININS else 'Monsieur'
        return f'{civilite} {nom_formate},'

    # Mode 2 : parser le string complet "XAVIER DENIS ALAIN BOREL (Président)"
    nom_clean = re.sub(r'\([^)]*\)', '', dirigeant_full)
    nom_clean = re.sub(r'\[[^\]]*\]', '', nom_clean).strip()
    if not nom_clean:
        return 'Madame, Monsieur,'

    parts = nom_clean.split()
    if not parts:
        return 'Madame, Monsieur,'

    if len(parts) < 2:
        return f'Monsieur {parts[0].title()},'

    # Le nom de famille est le DERNIER mot, le premier prenom est le PREMIER mot
    premier_prenom = parts[0]
    nom_famille = parts[-1]

    prenom_test = premier_prenom.split('-')[0].lower()
    civilite = 'Madame' if prenom_test in PRENOMS_FEMININS else 'Monsieur'

    return f'{civilite} {nom_famille.title()},'


def _get_secteur_text(prospect):
    """
    Retourne le texte du secteur d'activite pour le paragraphe d'intro.
    Cherche dans : libelle_naf, activite_declaree, code_naf.
    Fallback : 'dans son secteur d activite'.
    """
    # Essayer le libelle NAF (source principale)
    libelle = _clean(prospect.get('libelle_naf'))
    if not libelle:
        libelle = _clean(prospect.get('activite_declaree'))
    if not libelle:
        return ''

    # Minuscule pour insertion naturelle dans la phrase
    return libelle[0].lower() + libelle[1:]


def _download_logo(domain):
    """
    Telecharge le logo d'une entreprise. Essaie plusieurs sources.
    Retourne les bytes de l'image ou None.
    Seuil : > 1000 bytes pour eviter les favicons flous.
    """
    if not domain:
        return None

    sources = [
        f'https://img.companyenrich.com/logo?domain={domain}&format=png',
        f'https://logo.clearbit.com/{domain}',
        f'https://www.google.com/s2/favicons?domain={domain}&sz=128',
    ]

    for url in sources:
        try:
            r = requests.get(url, timeout=5)
            if r.status_code == 200 and len(r.content) > 1000:
                return r.content
        except Exception:
            continue

    return None


def _generate_intro_with_ai(entreprise, dirigeant, annee_creation, activite_naf, region, api_key):
    """
    Utilise Claude API pour generer un paragraphe d'intro naturel et personnalise.
    Retourne le texte ou None si pas de cle API ou erreur.
    """
    if not api_key:
        return None

    import anthropic
    client = anthropic.Anthropic(api_key=api_key)

    prompt = f"""Tu dois rédiger UN SEUL paragraphe pour une lettre de prospection bancaire (Banque Mirabaud).

Contexte :
- Entreprise : {entreprise}
- Dirigeant : {dirigeant or 'inconnu'}
- Année de création : {annee_creation or 'inconnue'}
- Code/libellé NAF : {activite_naf or 'inconnu'}
- Région : {region or 'France'}

Format EXACT à respecter (ne change PAS la structure, remplis juste les crochets) :
"Sous votre impulsion et depuis [année], [Nom entreprise propre] s'est affirmée comme une entreprise référente dans le secteur [reformule le secteur NAF de manière naturelle et professionnelle], et ce, dans un environnement en perpétuelle évolution ([3 défis/tendances actuels du secteur sous forme de mots-clefs séparés par des virgules])."

Règles :
- Le secteur doit être reformulé naturellement (PAS le libellé NAF brut)
  Exemple : "autres intermédiations monétaires" → "des services bancaires et financiers"
  Exemple : "travaux d'installation d'équipements thermiques" → "de l'installation et la maintenance d'équipements énergétiques"
- Les 3 défis doivent être pertinents pour CE secteur spécifique
  Exemple banque : "digitalisation des services, évolution réglementaire, enjeux ESG"
  Exemple construction : "transition bas carbone, industrialisation des chantiers, normes environnementales"
- Si le dirigeant est inconnu, commence par "Depuis [année]," sans "Sous votre impulsion"
- Si l'année est inconnue, commence directement par "[Nom entreprise] s'est affirmée..."
- Le nom de l'entreprise doit être PROPRE (pas de doublon, version courte/connue)
- UNE SEULE phrase, pas de retour à la ligne
- Ton professionnel, sobre, style banque d'affaires
- Maximum 250 caractères

Réponds UNIQUEMENT avec le paragraphe, sans guillemets, sans explication."""

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}]
        )
        text = response.content[0].text.strip()
        # Retirer les guillemets si Claude en a mis
        if text.startswith('"') and text.endswith('"'):
            text = text[1:-1]
        return text
    except Exception as e:
        print(f"  Claude API error: {e}")
        return None


# ---------------------------------------------------------------------------
# LetterGenerator
# ---------------------------------------------------------------------------

class LetterGenerator:
    """Genere des lettres Word a partir du template Mirabaud."""

    def __init__(self, template_path=None, output_dir="lettres", api_key=None):
        self.template_path = template_path or TEMPLATE_PATH
        self.output_dir = output_dir
        self.api_key = api_key or os.environ.get('ANTHROPIC_API_KEY', '')
        os.makedirs(output_dir, exist_ok=True)

    def generate_letter(self, prospect: dict) -> BytesIO:
        """
        Copie le template Mirabaud et personnalise :
        - Logo (P0) : remplace Bonna Sabla par le logo du prospect
        - Date (P2) : date du jour en francais
        - Salutation (P9) : civilite + nom de famille (format Titre)
        - Intro (P11) : nom entreprise, anciennete, secteur d'activite
        Retourne un BytesIO contenant le .docx.
        """
        doc = Document(self.template_path)

        # Extraire et nettoyer les donnees
        entreprise_raw = _clean(prospect.get('nom_entreprise'), 'votre entreprise')
        entreprise = _clean_company_name(entreprise_raw)
        dirigeant = _clean(prospect.get('dirigeant_enrichi')) or _clean(prospect.get('dirigeant_principal'))
        secteur = _get_secteur_text(prospect)
        date_creation = _clean(prospect.get('date_creation'))
        site_web = _clean(prospect.get('site_web'))
        region = _clean(prospect.get('region'))

        # Activite NAF pour le prompt IA
        activite_naf = _clean(prospect.get('libelle_naf'))
        if not activite_naf:
            activite_naf = _clean(prospect.get('code_naf'))

        # Anciennete
        annee_creation = ''
        if date_creation and len(date_creation) >= 4:
            try:
                annee_creation = str(int(date_creation[:4]))
            except ValueError:
                pass

        # --- Logo (P0) ---
        self._replace_logo(doc, site_web)

        # --- Date (P2) ---
        self._clear_and_set(doc.paragraphs[2], f"Paris, le {_format_date_fr()}")

        # --- Civilite (P9) : utiliser nom/prenom separes si disponibles ---
        dirigeant_nom = _clean(prospect.get('dirigeant_nom'))
        dirigeant_prenom = _clean(prospect.get('dirigeant_prenom'))
        salutation = _format_civilite(dirigeant, nom=dirigeant_nom, prenom=dirigeant_prenom)
        self._clear_and_set(doc.paragraphs[9], salutation)

        # --- Intro (P11) : IA si disponible, sinon fallback ---
        intro_ai = _generate_intro_with_ai(
            entreprise=entreprise,
            dirigeant=dirigeant,
            annee_creation=annee_creation,
            activite_naf=activite_naf,
            region=region,
            api_key=self.api_key,
        )

        if intro_ai:
            intro = intro_ai
        else:
            intro = self._build_intro(entreprise, annee_creation, secteur, dirigeant)

        self._clear_and_set(doc.paragraphs[11], intro)

        # Sauvegarder en memoire
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _clear_and_set(self, paragraph, text):
        """Vide toutes les runs d'un paragraphe puis ecrit dans la premiere."""
        for run in paragraph.runs:
            run.text = ''
        if paragraph.runs:
            paragraph.runs[0].text = text
        else:
            paragraph.add_run(text)

    def _replace_logo(self, doc, site_web):
        """
        Supprime le logo Bonna Sabla (drawing dans P0).
        Si le prospect a un site web, insere son logo a la place.
        """
        p0 = doc.paragraphs[0]
        ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # Supprimer tous les drawings de P0
        for run_elem in p0._element.findall(f'{{{ns_w}}}r'):
            for drawing in run_elem.findall(f'{{{ns_w}}}drawing'):
                run_elem.remove(drawing)

        # Telecharger et inserer le logo du prospect
        domain = _extract_domain(site_web)
        logo_bytes = _download_logo(domain)
        if logo_bytes:
            try:
                img_stream = BytesIO(logo_bytes)
                p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p0.runs[0] if p0.runs else p0.add_run()
                run.add_picture(img_stream, width=Inches(1.0))
            except Exception:
                pass

    def _build_intro(self, entreprise, annee_creation, secteur, dirigeant):
        """Construit le paragraphe d'introduction personnalise (P11). Fallback sans IA."""
        parts = []

        if dirigeant and annee_creation:
            parts.append(
                f"Sous votre impulsion et depuis {annee_creation}, "
                f"{entreprise} s\u2019est affirm\u00e9e comme une entreprise "
                f"r\u00e9f\u00e9rente"
            )
        elif annee_creation:
            parts.append(
                f"Depuis {annee_creation}, {entreprise} s\u2019est affirm\u00e9e "
                f"comme une entreprise r\u00e9f\u00e9rente"
            )
        else:
            parts.append(
                f"{entreprise} s\u2019est affirm\u00e9e comme une entreprise "
                f"r\u00e9f\u00e9rente"
            )

        if secteur:
            vowels = 'aeiouyéèêëàâäôöùûüîï'
            de = "d\u2019" if secteur[0].lower() in vowels else "de "
            parts.append(f" dans le secteur {de}{secteur}")
        else:
            parts.append(" dans son secteur d\u2019activit\u00e9")

        parts.append(
            ", et ce, dans un environnement en perp\u00e9tuelle \u00e9volution."
        )

        return ''.join(parts)

    def generate_filename(self, prospect: dict) -> str:
        """Genere un nom de fichier normalise pour la lettre."""
        nom = _clean(prospect.get('nom_entreprise'), 'prospect')
        nom = _clean_company_name(nom)
        safe_nom = re.sub(r'[^\w\s-]', '', nom)
        safe_nom = re.sub(r'\s+', '_', safe_nom.strip())[:50]
        return f"Lettre_{safe_nom}.docx"

    def generate_all(self, prospects_df) -> list:
        """
        Genere une lettre pour chaque prospect du DataFrame.
        Retourne la liste des chemins de fichiers crees.
        """
        os.makedirs(self.output_dir, exist_ok=True)
        filepaths = []

        for _, row in prospects_df.iterrows():
            prospect = row.to_dict()
            try:
                buf = self.generate_letter(prospect)
                filename = self.generate_filename(prospect)
                filepath = os.path.join(self.output_dir, filename)
                with open(filepath, 'wb') as f:
                    f.write(buf.getvalue())
                filepaths.append(filepath)
            except Exception as e:
                nom = prospect.get('nom_entreprise', '?')
                print(f"  Erreur lettre pour {nom}: {e}")

        return filepaths
